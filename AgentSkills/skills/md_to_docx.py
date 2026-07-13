"""
md_to_docx.py - Markdown 转 DOCX 转换工具

将 Markdown 格式的文件或文本转换为 Word (.docx) 文档，
支持标题、段落、粗体/斜体、代码块、列表、表格、图片、链接等元素。
"""

import os
import re
from datetime import datetime
from typing import Optional


def _ensure_dependencies():
    """确保所需的第三方库已安装"""
    missing = []
    try:
        import markdown
    except ImportError:
        missing.append("markdown")

    try:
        from docx import Document
    except ImportError:
        missing.append("python-docx")

    try:
        import bs4
    except ImportError:
        missing.append("beautifulsoup4")

    if missing:
        import subprocess
        deps = " ".join(missing)
        result = subprocess.run(
            f"pip install {deps}",
            shell=True, capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            return f"❌ 依赖安装失败: {result.stderr}"
        return f"✅ 已安装依赖: {', '.join(missing)}"
    return "✅ 依赖已就绪"


def convert_md_content_to_docx(
    md_content: str,
    output_path: str,
    title: Optional[str] = None,
    author: Optional[str] = None,
) -> str:
    """将 Markdown 文本内容转换为 DOCX 文件

    Args:
        md_content: Markdown 格式的文本内容
        output_path: 输出文件路径（自动补全 .docx 后缀）
        title: 文档标题（可选）
        author: 作者名（可选）

    Returns:
        成功或失败的消息
    """
    try:
        # 确保依赖已安装
        dep_result = _ensure_dependencies()

        # 确保路径以 .docx 结尾
        if not output_path.endswith(".docx"):
            output_path = output_path + ".docx"

        # 自动创建父目录
        parent_dir = os.path.dirname(output_path)
        if parent_dir and not os.path.exists(parent_dir):
            os.makedirs(parent_dir, exist_ok=True)

        # 导入所需库
        import markdown as md_lib
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from bs4 import BeautifulSoup

        # 将 Markdown 转换为 HTML
        html = md_lib.markdown(
            md_content,
            extensions=[
                'extra',         # 包含表格、围栏代码块、脚注等
                'codehilite',    # 代码高亮
                'toc',           # 目录
                'sane_lists',    # 智能列表
            ]
        )

        # 解析 HTML
        soup = BeautifulSoup(html, 'html.parser')

        # 创建 Word 文档
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # 添加文档标题（如果提供了 title）
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加元信息
        if author or title:
            meta_paragraph = doc.add_paragraph()
            meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if author:
                meta_run = meta_paragraph.add_run(f"作者: {author}  |  ")
            meta_run = meta_paragraph.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(128, 128, 128)
            doc.add_paragraph()  # 空行

        # 解析 HTML 元素并添加到文档
        for element in soup.children:
            if element.name is None:
                continue
            _parse_html_element(doc, element)

        # 保存文档
        doc.save(output_path)

        # 获取文件大小
        file_size = os.path.getsize(output_path)
        if file_size < 1024:
            size_str = f"{file_size} B"
        elif file_size < 1024 * 1024:
            size_str = f"{file_size / 1024:.1f} KB"
        else:
            size_str = f"{file_size / (1024 * 1024):.1f} MB"

        return f"✅ DOCX 文件已生成: {output_path} ({size_str})"

    except Exception as e:
        return f"❌ DOCX 文件生成失败: {e}"


def _parse_html_element(doc, element):
    """递归解析 HTML 元素并添加到 Document"""
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    tag = element.name
    if tag is None:
        return

    # ── 标题 (h1 ~ h6) ──
    if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag[1])
        heading = doc.add_heading(element.get_text(), level=level)
        return

    # ── 段落 ──
    if tag == 'p':
        para = doc.add_paragraph()
        _add_inline_elements(para, element)
        return

    # ── 代码块 ──
    if tag == 'pre':
        code_element = element.find('code')
        code_text = code_element.get_text() if code_element else element.get_text()
        # 使用等宽字体添加代码
        para = doc.add_paragraph()
        para.style = doc.styles['Normal']
        fmt = para.paragraph_format
        fmt.left_indent = Cm(1)
        # 添加浅灰色背景
        shading = para._p.get_or_add_pPr().makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'F2F2F2',
        })
        para._p.get_or_add_pPr().append(shading)
        run = para.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        return

    if tag == 'code':
        # 行内代码
        para = doc.add_paragraph()
        run = para.add_run(element.get_text())
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        return

    # ── 列表 ──
    if tag in ['ul', 'ol']:
        for li in element.find_all('li', recursive=False):
            para = doc.add_paragraph(style='List Bullet' if tag == 'ul' else 'List Number')
            _add_inline_elements(para, li)
        return

    # ── 表格 ──
    if tag == 'table':
        rows = element.find_all('tr')
        if not rows:
            return

        # 获取列数
        cols = 0
        for row in rows:
            cells = row.find_all(['th', 'td'])
            cols = max(cols, len(cells))

        if cols == 0:
            return

        # 创建 Word 表格
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j >= cols:
                    break
                word_cell = table.rows[i].cells[j]
                word_cell.text = cell.get_text().strip()

                # 表头加粗
                if cell.name == 'th':
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        return

    # ── 水平线 ──
    if tag == 'hr':
        doc.add_paragraph("─" * 50)
        return

    # ── 块引用 ──
    if tag == 'blockquote':
        for child in element.children:
            if child.name == 'p':
                para = doc.add_paragraph()
                para.style = doc.styles['Normal']
                fmt = para.paragraph_format
                fmt.left_indent = Cm(1.5)
                run = para.add_run(f"> {child.get_text()}")
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
        return

    # ── 图片 ──
    if tag == 'img':
        src = element.get('src', '')
        alt = element.get('alt', '图片')
        if src and os.path.exists(src):
            try:
                doc.add_picture(src, width=Inches(5))
            except Exception:
                para = doc.add_paragraph()
                run = para.add_run(f"[图片: {alt}]")
                run.italic = True
        else:
            para = doc.add_paragraph()
            run = para.add_run(f"[图片: {alt} - {src}]")
            run.italic = True
        return


def _add_inline_elements(para, parent_element):
    """处理段落中的内联元素（粗体、斜体、链接、行内代码等）"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    for child in parent_element.children:
        if child.name is None:
            # 纯文本节点
            if child.string and child.string.strip():
                para.add_run(child.string)
            continue

        if child.name == 'strong' or child.name == 'b':
            run = para.add_run(child.get_text())
            run.bold = True
            continue

        if child.name == 'em' or child.name == 'i':
            run = para.add_run(child.get_text())
            run.italic = True
            continue

        if child.name == 'code':
            run = para.add_run(child.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # 添加浅灰色背景
            shading = run._r.get_or_add_rPr().makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'F2F2F2',
            })
            run._r.get_or_add_rPr().append(shading)
            continue

        if child.name == 'a':
            href = child.get('href', '')
            text = child.get_text()
            if href:
                run = para.add_run(f"{text} ({href})")
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.underline = True
            else:
                para.add_run(text)
            continue

        if child.name == 'br':
            para.add_run("\n")
            continue

        if child.name == 'img':
            src = child.get('src', '')
            alt = child.get('alt', '图片')
            run = para.add_run(f"[图片: {alt}]")
            run.italic = True
            continue

        # 递归处理子元素
        _add_inline_elements(para, child)


def convert_md_to_docx(
    md_file_path: str,
    output_path: Optional[str] = None,
    title: Optional[str] = None,
    author: Optional[str] = None,
) -> str:
    """将 Markdown 文件转换为 DOCX 文件

    Args:
        md_file_path: Markdown 文件路径
        output_path: 输出文件路径（默认与输入文件同名，扩展名为 .docx）
        title: 文档标题（可选）
        author: 作者名（可选）

    Returns:
        成功或失败的消息
    """
    try:
        # 读取 Markdown 文件
        if not os.path.exists(md_file_path):
            return f"❌ 文件不存在: {md_file_path}"

        with open(md_file_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        # 默认输出路径
        if output_path is None:
            base = os.path.splitext(md_file_path)[0]
            output_path = base + ".docx"

        # 调用内容转换函数
        return convert_md_content_to_docx(
            md_content=md_content,
            output_path=output_path,
            title=title,
            author=author,
        )

    except Exception as e:
        return f"❌ DOCX 转换失败: {e}"
